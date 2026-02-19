# -*- coding: utf-8 -*-
"""
Transcribe PDFs to DOCX, removing ruby (furigana).

Requirements (already present in this repo environment):
  - PyMuPDF (fitz)
  - python-docx

Key behaviors:
  - 1 PDF -> 1 DOCX
  - Each PDF page starts on a new DOCX page (page break between pages)
  - Ruby (furigana) is detected by layout: small kana text positioned just above
    larger base text with horizontal overlap, then removed.

Usage:
  python tools/pdf_transcribe_to_docx.py ^
    --input-root "C:\\path\\to\\pdf_root" ^
    --output-root "C:\\path\\to\\pdf_root\\docx_out"
"""

from __future__ import annotations

import argparse
import re
import sys
import unicodedata
from dataclasses import dataclass
from pathlib import Path
from typing import Iterable

import fitz  # PyMuPDF
from docx import Document  # type: ignore
from docx.oxml.ns import qn  # type: ignore
from docx.shared import Pt  # type: ignore


@dataclass(frozen=True)
class TextSpan:
    span_id: int
    text: str
    size: float
    x0: float
    y0: float
    x1: float
    y1: float
    block_id: int
    line_id: int

    @property
    def width(self) -> float:
        return max(0.0, self.x1 - self.x0)

    @property
    def height(self) -> float:
        return max(0.0, self.y1 - self.y0)


@dataclass(frozen=True)
class TextLine:
    block_id: int
    line_id: int
    x0: float
    y0: float
    x1: float
    y1: float
    span_ids: tuple[int, ...]


_KANA_RE = re.compile(r"[\u3040-\u309F\u30A0-\u30FF\u31F0-\u31FF\uFF65-\uFF9F]")
_KANJI_RE = re.compile(r"[\u4E00-\u9FFF]")
_CJK_RE_CLASS = r"\u3040-\u309F\u30A0-\u30FF\u4E00-\u9FFF"

_PAGE_NUM_RE = re.compile(r"^[-‐‑‒–—―ーｰ]{1,4}\s*\d+\s*[-‐‑‒–—―ーｰ]{1,4}$")
_KANA_LABEL_ONLY_RE = re.compile(
    r"^(?:[（(][イロハニホヘト][）)])(?:\s+[（(][イロハニホヘト][）)])+$"
)


def _has_kana(text: str) -> bool:
    return bool(_KANA_RE.search(text))


def _has_kanji(text: str) -> bool:
    return bool(_KANJI_RE.search(text))


def _overlap_ratio(a: TextSpan, b: TextSpan) -> float:
    overlap = min(a.x1, b.x1) - max(a.x0, b.x0)
    if overlap <= 0:
        return 0.0
    denom = min(a.width, b.width)
    if denom <= 0:
        return 0.0
    return overlap / denom


def _detect_ruby_span_ids(spans: list[TextSpan]) -> set[int]:
    """
    Detect ruby spans by geometry.

    Heuristic:
      - span contains kana (ruby is typically kana)
      - there is a larger span immediately below it, with horizontal overlap,
        and very small vertical gap.
    """
    # Sort by y0 for searching base text below.
    spans_by_y = sorted(spans, key=lambda s: (s.y0, s.x0))
    y0_list = [s.y0 for s in spans_by_y]

    # Simple binary search helpers (avoid importing bisect in tight loop)
    def lower_bound(value: float) -> int:
        lo = 0
        hi = len(y0_list)
        while lo < hi:
            mid = (lo + hi) // 2
            if y0_list[mid] < value:
                lo = mid + 1
            else:
                hi = mid
        return lo

    ruby_ids: set[int] = set()

    for s in spans:
        t = s.text.strip()
        if not t:
            continue
        # Ruby is almost always kana; this prevents accidental removal.
        if not _has_kana(t):
            continue
        # Ruby usually does not include kanji; allow mixed cases but penalize.
        if _has_kanji(t):
            continue

        # Search a narrow band just below the span.
        search_y_min = s.y0
        # Ruby-to-base gap is typically within ~1-2 ruby heights.
        search_y_max = s.y1 + max(8.0, s.height * 2.2)

        start = lower_bound(search_y_min)
        best_base: TextSpan | None = None
        best_score: float | None = None

        for b in spans_by_y[start:]:
            if b.span_id == s.span_id:
                continue
            if b.y0 > search_y_max:
                break

            # Base must be below or slightly overlapping.
            gap = b.y0 - s.y1
            if gap < -max(1.0, s.height * 0.35):
                continue

            # Base should be larger.
            if b.size < max(s.size * 1.18, s.size + 1.0):
                continue

            # Must overlap horizontally.
            ov = _overlap_ratio(s, b)
            if ov < 0.35:
                continue

            # Prefer closest vertical match and strong overlap.
            # Smaller score is better.
            score = abs(gap) - (ov * 0.5)
            if best_score is None or score < best_score:
                best_score = score
                best_base = b

        if best_base is not None:
            ruby_ids.add(s.span_id)

    return ruby_ids


def _join_spans_into_line(spans: list[TextSpan]) -> str:
    spans_sorted = sorted(spans, key=lambda s: (s.x0, s.y0))
    out: list[str] = []
    prev: TextSpan | None = None
    for s in spans_sorted:
        text = s.text.replace("\n", " ")
        if not text:
            continue
        if prev is None:
            out.append(text)
            prev = s
            continue

        gap = s.x0 - prev.x1
        # Insert a space if the visual gap is meaningful.
        if gap > max(2.0, max(prev.size, s.size) * 0.35):
            out.append(" ")
        out.append(text)
        prev = s

    return "".join(out)


def _normalize_text(text: str, *, reiwa_year: int | None) -> str:
    # Fix common symbol mis-decoding in these PDFs.
    # "姶" frequently appears where "×" is intended.
    text = text.replace("姶", "×")

    # Normalize width variants etc.
    text = unicodedata.normalize("NFKC", text)

    # These PDFs often place each glyph as a separate text object, so our visual
    # join may introduce spaces between Japanese glyphs. Remove those.
    text = re.sub(rf"(?<=[{_CJK_RE_CLASS}])\s+(?=[{_CJK_RE_CLASS}])", "", text)
    text = re.sub(r"(?<=\d)\s+(?=\d)", "", text)
    text = re.sub(rf"(?<=\d)\s+(?=[{_CJK_RE_CLASS}])", "", text)
    text = re.sub(rf"(?<=[{_CJK_RE_CLASS}])\s+(?=\d)", "", text)

    # Remove spaces around control characters to make pattern replacement easier.
    text = re.sub(r"\s*([\x00-\x1F])\s*", r"\1", text)

    # Replace known "missing digit" patterns in these PDFs.
    if reiwa_year is not None:
        # Example: "令和\x02 年度" -> "令和3年度" (digit inferred from filename)
        text = re.sub(r"令和[\x00-\x1F]+年度", f"令和{reiwa_year}年度", text)

    # The exam name is always "1級土木施工管理技術検定" in this dataset.
    text = re.sub(r"[\x00-\x1F]+級土木施工管理技術検定", "1級土木施工管理技術検定", text)

    # Drop remaining control characters.
    text = re.sub(r"[\x00-\x08\x0B-\x1F]", "", text)

    # Drop private-use glyphs used for decorative lines.
    text = re.sub(r"[\uE000-\uF8FF]+", "", text)

    return text.strip()


def _should_skip_line(text: str) -> bool:
    s = text.strip()
    if not s:
        return True
    # Page number like "-1-" or "― 1 ―"
    if _PAGE_NUM_RE.match(s):
        return True
    # This caption/title is shown as an image on the website, so do not transcribe it.
    if s == "土の構成の模式図":
        return True
    # Column labels line like "（イ） （ロ） （ハ）" (we inline labels into choices instead).
    if _KANA_LABEL_ONLY_RE.match(s):
        return True
    return False


def _fix_choice_placeholders(line: str) -> str:
    """
    Fix specific placeholders used in some questions.

    Example (R7 No.2):
      - "B,C,D" should be displayed as kana labels: "（イ）,（ロ）,（ハ）"
      - "B" / "C" / "D" inside the explanation should become "（イ）" / "（ロ）" / "（ハ）"
      - Header line "B C D" should become "（イ） （ロ） （ハ）" (later skipped)

    This MUST NOT touch formula variables like "D" in other questions, so the
    replacement is intentionally narrow (only lines that clearly refer to blanks).
    """
    s = line.strip()
    if not s:
        return line

    # "B C D" header line above choices.
    if re.fullmatch(r"(?:B\s*C\s*D|B\s+C\s+D)", s):
        return "（イ） （ロ） （ハ）"

    # Question stem line: "...文章中のB,C,Dに当て..."
    if "文章中" in s and ("当て" in s or "当" in s) and re.search(r"B\s*,\s*C\s*,\s*D", s):
        fixed = s.replace("文章中のの", "文章中の")
        fixed = re.sub(r"B\s*,\s*C\s*,\s*D", "（イ）,（ロ）,（ハ）", fixed)
        return fixed

    # Explanation lines: contain "含水比" and "という" and placeholders.
    if ("含水比" in s or "語句" in s) and "という" in s and re.search(r"(?<![A-Za-z0-9])[BCD](?![A-Za-z0-9])", s):
        fixed = re.sub(r"(?<![A-Za-z0-9])B(?![A-Za-z0-9])", "（イ）", s)
        fixed = re.sub(r"(?<![A-Za-z0-9])C(?![A-Za-z0-9])", "（ロ）", fixed)
        fixed = re.sub(r"(?<![A-Za-z0-9])D(?![A-Za-z0-9])", "（ハ）", fixed)
        # Remove spaces introduced around the blanks.
        fixed = (
            fixed.replace("・ （", "・（")
            .replace("と （", "と（")
            .replace("） と", "）と")
            .replace("） を", "）を")
            .replace("） の", "）の")
            .replace(", （", ",（")
            .replace("， （", "，（")
        )
        return fixed

    # Special combined sentence: "B と C の含水比の幅を, D という。"
    if "含水比" in s and "幅" in s and "という" in s and ("B" in s or "C" in s or "D" in s):
        fixed = s
        fixed = re.sub(r"(?<![A-Za-z0-9])B(?![A-Za-z0-9])", "（イ）", fixed)
        fixed = re.sub(r"(?<![A-Za-z0-9])C(?![A-Za-z0-9])", "（ロ）", fixed)
        fixed = re.sub(r"(?<![A-Za-z0-9])D(?![A-Za-z0-9])", "（ハ）", fixed)
        fixed = (
            fixed.replace("・ （", "・（")
            .replace("と （", "と（")
            .replace("） と", "）と")
            .replace("） の", "）の")
            .replace(", （", ",（")
            .replace("， （", "，（")
        )
        return fixed

    return line


_CHOICE_LABELS = ["(イ)", "(ロ)", "(ハ)", "(ニ)", "(ホ)", "(ヘ)", "(ト)"]
_CHOICE_SPLIT_RE = re.compile(r"\s*(?:\.{3}|…|⋯)\s*")


def _rewrite_choice_line_with_kana_labels(line: str) -> str:
    """
    Rewrite choice lines like:
      (1) A ... B ... C
    into:
      (1) (イ)A ... (ロ)B ... (ハ)C

    Applies to other similar multi-column choices as well.
    """
    m = re.match(r"^\((\d+)\)\s*(.+)$", line.strip())
    if not m:
        return line

    num = m.group(1)
    rest = m.group(2).strip()

    # Already labeled.
    if any(lbl in rest for lbl in _CHOICE_LABELS[:3]):
        return line

    parts = [p.strip() for p in _CHOICE_SPLIT_RE.split(rest)]
    parts = [p for p in parts if p]
    if len(parts) < 2:
        return line
    if len(parts) > len(_CHOICE_LABELS):
        return line

    labeled = [f"{_CHOICE_LABELS[i]}{parts[i]}" for i in range(len(parts))]
    return f"({num}) " + " ... ".join(labeled)


def _infer_reiwa_year_from_filename(pdf_path: Path) -> int | None:
    """
    Infer Reiwa year from filenames like "R3gakkaA_mondai.pdf".
    Returns None when unknown.
    """
    m = re.match(r"(?i)R([1-9]\d?)", pdf_path.stem)
    if not m:
        return None
    try:
        return int(m.group(1))
    except ValueError:
        return None


def extract_page_text_lines(page: fitz.Page, *, reiwa_year: int | None) -> list[str]:
    d = page.get_text("dict", sort=True)

    spans: list[TextSpan] = []
    next_span_id = 1

    for b_idx, block in enumerate(d.get("blocks", [])):
        if block.get("type") != 0:
            continue
        for l_idx, line in enumerate(block.get("lines", [])):
            for span in line.get("spans", []):
                txt = span.get("text", "")
                if not txt or txt.isspace():
                    continue
                x0, y0, x1, y1 = span.get("bbox", (0.0, 0.0, 0.0, 0.0))
                size = float(span.get("size", 0.0))
                spans.append(
                    TextSpan(
                        span_id=next_span_id,
                        text=txt,
                        size=size,
                        x0=float(x0),
                        y0=float(y0),
                        x1=float(x1),
                        y1=float(y1),
                        block_id=b_idx,
                        line_id=l_idx,
                    )
                )
                next_span_id += 1

    ruby_ids = _detect_ruby_span_ids(spans)
    kept_spans = [s for s in spans if s.span_id not in ruby_ids]

    # Cluster into visual lines ourselves (these PDFs have each glyph as its own block).
    lines: list[list[TextSpan]] = _cluster_spans_into_lines(kept_spans)

    out_lines: list[str] = []
    for line_spans in lines:
        line_text = _join_spans_into_line(line_spans)
        line_text = _normalize_text(line_text, reiwa_year=reiwa_year)
        if not line_text:
            continue
        line_text = _fix_choice_placeholders(line_text)
        line_text = _rewrite_choice_line_with_kana_labels(line_text)
        if _should_skip_line(line_text):
            continue
        out_lines.append(line_text)

    return out_lines


def _cluster_spans_into_lines(spans: list[TextSpan]) -> list[list[TextSpan]]:
    if not spans:
        return []

    spans_sorted = sorted(spans, key=lambda s: (s.y0, s.x0))

    lines: list[list[TextSpan]] = []
    cur: list[TextSpan] = []
    cur_yc: float | None = None
    cur_size: float = 0.0

    def tol_for(size: float) -> float:
        # Tolerance for grouping spans into the same line.
        return max(2.0, size * 0.35)

    for s in spans_sorted:
        yc = (s.y0 + s.y1) / 2.0
        if not cur:
            cur = [s]
            cur_yc = yc
            cur_size = s.size
            continue

        assert cur_yc is not None
        tol = tol_for(cur_size)
        if abs(yc - cur_yc) <= tol:
            cur.append(s)
            # Update running averages (stable enough for this task).
            cur_yc = (cur_yc * 0.8) + (yc * 0.2)
            cur_size = (cur_size * 0.8) + (s.size * 0.2)
        else:
            lines.append(cur)
            cur = [s]
            cur_yc = yc
            cur_size = s.size

    if cur:
        lines.append(cur)

    # Order lines by their top y then left x.
    lines.sort(key=lambda ln: (min(s.y0 for s in ln), min(s.x0 for s in ln)))
    return lines


def _configure_doc_defaults(doc: Document) -> None:
    # Make paragraphs compact (no extra spacing) so page boundaries are closer to PDF.
    style = doc.styles["Normal"]
    style.font.name = "Yu Mincho"
    style.font.size = Pt(10.5)
    # Ensure East Asia font is set (Word uses a separate slot for Japanese fonts).
    try:
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Yu Mincho")  # type: ignore[attr-defined]
    except Exception:
        pass
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.line_spacing = 1.0


def transcribe_pdf_to_docx(pdf_path: Path, output_docx_path: Path) -> None:
    reiwa_year = _infer_reiwa_year_from_filename(pdf_path)

    doc = Document()
    _configure_doc_defaults(doc)

    pdf = fitz.open(pdf_path)
    try:
        for i, page in enumerate(pdf):
            lines = extract_page_text_lines(page, reiwa_year=reiwa_year)
            if not lines:
                # Preserve page boundary even when text is empty.
                doc.add_paragraph("")
            else:
                for line in lines:
                    doc.add_paragraph(line)
            if i != len(pdf) - 1:
                doc.add_page_break()
    finally:
        pdf.close()

    output_docx_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_docx_path))


def _iter_pdf_files(root: Path) -> Iterable[Path]:
    # Deterministic ordering.
    return sorted(root.rglob("*.pdf"), key=lambda p: str(p).lower())


def main(argv: list[str]) -> int:
    parser = argparse.ArgumentParser(description="Transcribe PDFs to DOCX (remove ruby).")
    parser.add_argument("--input-root", type=Path, required=True, help="Root folder to search PDFs under.")
    parser.add_argument("--output-root", type=Path, required=True, help="Root folder to write DOCX files to.")
    parser.add_argument(
        "--only",
        type=Path,
        default=None,
        help="Process only this PDF (absolute path or path relative to --input-root).",
    )
    parser.add_argument("--skip-existing", action="store_true", help="Skip if output docx already exists.")
    args = parser.parse_args(argv)

    input_root: Path = args.input_root
    output_root: Path = args.output_root

    if not input_root.exists():
        print(f"ERROR: input root not found: {input_root}", file=sys.stderr)
        return 2

    if args.only:
        only_path = args.only
        if not only_path.is_absolute():
            only_path = input_root / only_path
        if not only_path.exists():
            print(f"ERROR: --only not found: {only_path}", file=sys.stderr)
            return 2
        if only_path.suffix.lower() != ".pdf":
            print(f"ERROR: --only must be a .pdf file: {only_path}", file=sys.stderr)
            return 2
        pdf_paths = [only_path]
    else:
        pdf_paths = list(_iter_pdf_files(input_root))
    if not pdf_paths:
        print(f"ERROR: no PDFs found under: {input_root}", file=sys.stderr)
        return 2

    ok = 0
    skipped = 0
    failed = 0

    for pdf_path in pdf_paths:
        rel = pdf_path.relative_to(input_root)
        out_path = (output_root / rel).with_suffix(".docx")
        if args.skip_existing and out_path.exists():
            skipped += 1
            continue
        try:
            transcribe_pdf_to_docx(pdf_path, out_path)
            ok += 1
        except Exception as e:
            failed += 1
            print(f"ERROR: failed: {pdf_path} -> {out_path}: {e}", file=sys.stderr)

    print(f"Done. OK={ok} Skipped={skipped} Failed={failed}")
    return 0 if failed == 0 else 1


if __name__ == "__main__":
    raise SystemExit(main(sys.argv[1:]))
